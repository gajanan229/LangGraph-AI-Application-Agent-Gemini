import unittest
from unittest.mock import patch, MagicMock, mock_open
import tempfile
import os
from docx import Document
from core.doc_generator import create_resume_pdf, create_cover_letter_pdf

# Constants for template paths from the user-provided structure
RESUME_TEMPLATE_PATH = "Templates/resume_template.docx"
PROJECT_TEMPLATE_PATH = "Templates/Project_template.docx"
COVER_LETTER_TEMPLATE_PATH = "Templates/cover_letter_template.docx"


class TestDocGenerator(unittest.TestCase):
    """Unit tests for the final document generation module."""

    @patch("core.doc_generator.convert")
    @patch("core.doc_generator.Document")
    @patch("builtins.open", new_callable=mock_open, read_data=b"dummy pdf content")
    def test_create_resume_pdf_returns_bytes(self, mock_file, mock_doc_class, mock_convert):
        """
        Tests that create_resume_pdf correctly fills templates and returns
        a non-empty byte string, mocking the actual PDF conversion.
        """
        # Arrange
        summary = "This is a test summary."
        projects = [
            {"title": "Project Apollo", "rewritten_text": "Bullet point 1\nBullet point 2"}
        ]
        
        # Mock Document and its methods
        mock_doc = MagicMock()
        mock_doc.paragraphs = []  # Empty paragraphs for simplicity
        mock_doc.save = MagicMock()
        mock_doc_class.return_value = mock_doc

        # Act
        pdf_bytes = create_resume_pdf(
            RESUME_TEMPLATE_PATH, PROJECT_TEMPLATE_PATH, summary, projects
        )

        # Assert
        self.assertIsInstance(pdf_bytes, bytes)
        self.assertGreater(len(pdf_bytes), 0)
        self.assertEqual(pdf_bytes, b"dummy pdf content")
        mock_convert.assert_called_once()
        mock_doc.save.assert_called_once()

    @patch("core.doc_generator.convert")
    @patch("core.doc_generator.Document")
    @patch("builtins.open", new_callable=mock_open, read_data=b"dummy cl content")
    def test_create_cover_letter_pdf_returns_bytes(self, mock_file, mock_doc_class, mock_convert):
        """
        Tests that create_cover_letter_pdf fills placeholders and returns
        a non-empty byte string, mocking the PDF conversion.
        """
        # Arrange
        intro = "Dear Hiring Manager,"
        body = "This is the main body of the cover letter."
        conclusion = "Sincerely, John Doe"
        
        # Mock Document and its methods
        mock_doc = MagicMock()
        mock_doc.paragraphs = []  # Empty paragraphs for simplicity
        mock_doc.save = MagicMock()
        mock_doc_class.return_value = mock_doc

        # Act
        pdf_bytes = create_cover_letter_pdf(
            COVER_LETTER_TEMPLATE_PATH, intro, body, conclusion
        )

        # Assert
        self.assertIsInstance(pdf_bytes, bytes)
        self.assertGreater(len(pdf_bytes), 0)
        self.assertEqual(pdf_bytes, b"dummy cl content")
        mock_convert.assert_called_once()
        mock_doc.save.assert_called_once()

    @patch("core.doc_generator.convert")
    @patch("core.doc_generator.tempfile.TemporaryDirectory")
    def test_placeholder_replacement_in_resume(self, mock_temp_dir, mock_convert):
        """
        Verifies that summary placeholders are replaced and projects are manually inserted
        with proper formatting in the DOCX object before conversion.
        """
        # Arrange
        summary = "UNIQUE_SUMMARY_TEXT_123"
        project_title = "UNIQUE_PROJECT_TITLE_456"
        project_bullets = "UNIQUE_BULLET_POINT_789\nANOTHER_BULLET_POINT_101"
        
        projects = [
            {"title": project_title, "rewritten_text": project_bullets}
        ]

        # Create a temporary directory for testing
        temp_dir = tempfile.mkdtemp()
        mock_temp_dir.return_value.__enter__.return_value = temp_dir
        
        # Mock the file operations
        docx_path = os.path.join(temp_dir, "final_resume.docx")
        pdf_path = os.path.join(temp_dir, "final_resume.pdf")
        
        # Create a dummy PDF file for reading
        with open(pdf_path, "wb") as f:
            f.write(b"test pdf content")

        # Act
        pdf_bytes = create_resume_pdf(RESUME_TEMPLATE_PATH, PROJECT_TEMPLATE_PATH, summary, projects)
        
        # Assert
        self.assertTrue(mock_convert.called)
        
        # Verify the DOCX file was created and contains our content
        self.assertTrue(os.path.exists(docx_path))
        
        # Load the created DOCX and check content
        doc = Document(docx_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check that summary placeholder was replaced
        self.assertIn(summary, full_text)
        self.assertNotIn("[SUMMARY]", full_text)
        
        # Check that projects were manually inserted
        self.assertIn(project_title, full_text)
        self.assertIn("UNIQUE_BULLET_POINT_789", full_text)
        self.assertIn("ANOTHER_BULLET_POINT_101", full_text)
        
        # Cleanup
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(docx_path):
            os.remove(docx_path)
        os.rmdir(temp_dir)

    @patch("core.doc_generator.convert")
    @patch("core.doc_generator.tempfile.TemporaryDirectory")
    def test_cover_letter_placeholder_replacement(self, mock_temp_dir, mock_convert):
        """
        Verifies that cover letter placeholders are properly replaced.
        """
        # Arrange
        intro = "UNIQUE_INTRO_TEXT_123"
        body = "UNIQUE_BODY_TEXT_456\nSECOND_LINE_789"
        conclusion = "UNIQUE_CONCLUSION_TEXT_101"

        # Create a temporary directory for testing
        temp_dir = tempfile.mkdtemp()
        mock_temp_dir.return_value.__enter__.return_value = temp_dir
        
        # Mock the file operations
        docx_path = os.path.join(temp_dir, "final_cl.docx")
        pdf_path = os.path.join(temp_dir, "final_cl.pdf")
        
        # Create a dummy PDF file for reading
        with open(pdf_path, "wb") as f:
            f.write(b"test cl pdf content")

        # Act
        pdf_bytes = create_cover_letter_pdf(COVER_LETTER_TEMPLATE_PATH, intro, body, conclusion)
        
        # Assert
        self.assertTrue(mock_convert.called)
        
        # Verify the DOCX file was created and contains our content
        self.assertTrue(os.path.exists(docx_path))
        
        # Load the created DOCX and check content
        doc = Document(docx_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check that placeholders were replaced
        self.assertIn(intro, full_text)
        self.assertIn("UNIQUE_BODY_TEXT_456", full_text)
        self.assertIn("SECOND_LINE_789", full_text)
        self.assertIn(conclusion, full_text)
        
        # Check that placeholders are gone
        self.assertNotIn("[INTRODUCTION]", full_text)
        self.assertNotIn("[BODY]", full_text)
        self.assertNotIn("[CONCLUSION]", full_text)
        
        # Cleanup
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(docx_path):
            os.remove(docx_path)
        os.rmdir(temp_dir)


if __name__ == "__main__":
    unittest.main() 