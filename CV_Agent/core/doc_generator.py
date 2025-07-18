import io
import os
import tempfile
from typing import Dict, List, Optional
import pythoncom
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert


def _delete_paragraph(paragraph: Paragraph):
    """Safely removes a paragraph from a document."""
    p_element = paragraph._element
    if p_element.getparent() is not None:
        p_element.getparent().remove(p_element)


def _add_formatted_text_to_paragraph(p: Paragraph, text: str):
    """
    Adds text to a paragraph, parsing double asterisks for bolding.
    Example: "This is **bold** text."
    """
    p.clear()  # Clear any existing runs, like the placeholder text
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        # Parts at odd indices are bold
        is_bold = (i % 2 == 1)
        run = p.add_run(part)
        run.bold = is_bold
        # Preserve the font size of the original paragraph style if possible
        if p.style and p.style.font:
            run.font.size = p.style.font.size
        else: # Fallback for project bullets which might not have a style
            run.font.size = Pt(10.5)


def _find_first_paragraph_with_text(doc: DocxDocument, text: str) -> Optional[Paragraph]:
    """Finds the first paragraph containing the given text."""
    for p in doc.paragraphs:
        if text in p.text:
            return p
    return None


def _replace_text_in_paragraph(p: Paragraph, placeholder: str, value: str):
    """
    Replaces placeholder text in a paragraph while preserving formatting.
    Assumes the placeholder exists entirely within a single run.
    Ensures Times New Roman size 11 font is applied.
    """
    for run in p.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            # Ensure consistent font formatting
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            break  # Assume placeholder appears only once per paragraph


def _ensure_cover_letter_font_formatting(doc: DocxDocument):
    """
    Ensures all paragraphs in the cover letter use Times New Roman size 11.
    This is a fallback to maintain consistent formatting.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name != 'Times New Roman' or run.font.size != Pt(11):
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)


def _find_projects_insertion_point(doc: DocxDocument) -> Optional[Paragraph]:
    """
    Finds the appropriate insertion point for projects.
    First tries to find "[PROJECTS_SECTION]" placeholder, 
    then falls back to inserting before the last paragraph.
    """
    # First try to find a projects section placeholder
    projects_placeholder = _find_first_paragraph_with_text(doc, "[PROJECTS_SECTION]")
    if projects_placeholder:
        return projects_placeholder
    
    # Fallback: insert before the last paragraph (usually contact info or footer)
    if doc.paragraphs:
        return doc.paragraphs[-1]
    
    return None


def _create_project_title(doc: DocxDocument, title: str, insert_before: Paragraph) -> Paragraph:
    """
    Creates a properly formatted project title paragraph.
    Bold text, 10.5 font size.
    """
    new_para = insert_before.insert_paragraph_before(title)
    
    # Clear any existing runs and create a new one with formatting
    new_para.clear()
    run = new_para.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    
    # Remove extra spacing
    new_para.paragraph_format.space_before = Pt(0)
    new_para.paragraph_format.space_after = Pt(0)
    
    return new_para


def _create_project_bullet(doc: DocxDocument, bullet_text: str, insert_before: Paragraph) -> Paragraph:
    """
    Creates a properly formatted bullet point paragraph with bolding support.
    """
    new_para = insert_before.insert_paragraph_before() # Create empty para
    
    # Set bullet list style first
    try:
        new_para.style = 'List Paragraph'
    except Exception:
        # Fallback if 'List Bullet' style doesn't exist. Prepend bullet manually.
        print("List Bullet style doesn't exist. Prepending bullet manually.")
        bullet_text = f"• {bullet_text}" 

    # Now add the formatted text
    _add_formatted_text_to_paragraph(new_para, bullet_text)
    
    # Remove extra spacing
    new_para.paragraph_format.space_before = Pt(0)
    new_para.paragraph_format.space_after = Pt(0)
    
    return new_para


def create_resume_pdf(
    resume_template_path: str,
    project_template_path: str, # This will be ignored but kept for compatibility
    summary_text: str,
    project_data: List[Dict[str, str]],
) -> bytes:
    """
    Builds a resume by populating a DOCX template to preserve formatting,
    then converts it to PDF.
    """
    doc = Document(resume_template_path)

    # 1. Replace Summary
    summary_anchor = _find_first_paragraph_with_text(doc, "[SUMMARY]")
    if summary_anchor:
        _add_formatted_text_to_paragraph(summary_anchor, summary_text)

    # 2. Handle legacy project placeholders (remove if they exist)
    title_anchor = _find_first_paragraph_with_text(doc, "[PROJECT TITLE]")
    bullets_anchor = _find_first_paragraph_with_text(doc, "[PROJECT BULLET POINTS]")
    
    if title_anchor:
        _delete_paragraph(title_anchor)
    if bullets_anchor:
        _delete_paragraph(bullets_anchor)

    # 3. Insert Projects with Manual Formatting
    insertion_point = _find_projects_insertion_point(doc)
    
    if insertion_point and project_data:
        # Insert projects in reverse order since we're inserting before
        for project in reversed(project_data):
            # Get bullet points, filtering out empty lines
            bullet_points = [b.strip() for b in project.get("rewritten_text", "").split('\n') if b.strip()]
            
            # Since we're using insert_paragraph_before(), the last thing inserted appears first
            # So we insert title LAST so it appears ABOVE the bullet points
            # Titles are not expected to be bolded, so this is fine.
            _create_project_title(doc, project.get("title", ""), insertion_point)
            
            # Insert bullet points in normal order (not reversed)
            # The _create_project_bullet function now handles bolding
            for point in bullet_points:
                _create_project_bullet(doc, point, insertion_point)
        
        # Remove the placeholder if it was "[PROJECTS_SECTION]"
        if "[PROJECTS_SECTION]" in insertion_point.text:
            _delete_paragraph(insertion_point)

    # 4. Save and Convert to PDF
    with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmpdir:
        docx_path = os.path.join(tmpdir, "final_resume.docx")
        pdf_path = os.path.join(tmpdir, "final_resume.pdf")
        doc.save(docx_path)
        try:
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()  # Ensure COM is uninitialized
            
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes


def create_cover_letter_pdf(
    template_path: str, intro: str, body: str, conclusion: str
) -> bytes:
    """
    Builds a cover letter by populating a DOCX template, then converts to PDF.
    """
    doc = Document(template_path)

    # --- Find all placeholders first to get stable references ---
    intro_anchor = _find_first_paragraph_with_text(doc, "[INTRODUCTION]")
    body_anchor = _find_first_paragraph_with_text(doc, "[BODY]")
    conclusion_anchor = _find_first_paragraph_with_text(doc, "[CONCLUSION]")

    # --- Process Introduction ---
    if intro_anchor:
        _replace_text_in_paragraph(intro_anchor, "[INTRODUCTION]", intro)
        intro_anchor.paragraph_format.space_before = Pt(6)
        intro_anchor.paragraph_format.space_after = Pt(6)

    # --- Process Conclusion ---
    if conclusion_anchor:
        _replace_text_in_paragraph(conclusion_anchor, "[CONCLUSION]", conclusion)
        conclusion_anchor.paragraph_format.space_before = Pt(0)
        conclusion_anchor.paragraph_format.space_after = Pt(6)

    # --- Process Body (only if body content is provided) ---
    if body_anchor and body and body.strip():  # Only process if body has actual content
        body_lines = [line.strip() for line in body.split('\n') if line.strip()]
        if body_lines:
            # Replace the anchor paragraph's text with the first line.
            body_anchor.text = body_lines[0]
            # Ensure font formatting for the body anchor
            for run in body_anchor.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            body_anchor.paragraph_format.space_before = Pt(0)
            body_anchor.paragraph_format.space_after = Pt(6)

            # Find the paragraph following the anchor to use as a reference.
            anchor_element = body_anchor._p
            all_paragraphs = list(doc.paragraphs)
            reference_paragraph = None
            for i, p in enumerate(all_paragraphs):
                if p._p == anchor_element and i + 1 < len(all_paragraphs):
                    reference_paragraph = all_paragraphs[i + 1]
                    break
            
            # Insert the rest of the body lines before the reference paragraph.
            if reference_paragraph and len(body_lines) > 1:
                for line in reversed(body_lines[1:]):
                    new_para = reference_paragraph.insert_paragraph_before(line)
                    new_para.style = body_anchor.style
                    # Ensure font formatting for new paragraphs
                    for run in new_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    new_para.paragraph_format.space_before = Pt(0)
                    new_para.paragraph_format.space_after = Pt(6)
            elif len(body_lines) > 1:
                # Fallback if [BODY] is the last thing in the document.
                for line in body_lines[1:]:
                    new_para = doc.add_paragraph(line, style=body_anchor.style)
                    # Ensure font formatting for new paragraphs
                    for run in new_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    new_para.paragraph_format.space_before = Pt(0)
                    new_para.paragraph_format.space_after = Pt(6)
    # If body is empty or None, leave the [BODY] placeholder untouched to preserve formatting

    # --- Apply font formatting fallback to entire document ---
    _ensure_cover_letter_font_formatting(doc)

    # Save and Convert to PDF
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "final_cl.docx")
        pdf_path = os.path.join(tmpdir, "final_cl.pdf")
        doc.save(docx_path)
        try:
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()  # Ensure COM is uninitialized

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes 